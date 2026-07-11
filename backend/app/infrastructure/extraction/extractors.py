"""Per-format text extractors.

Each extractor handles one canonical MIME type (as validated and persisted by
the documents feature) and is registered in the provider registry under
``kind="extractor"``. Parsing is CPU-bound library work, so it is offloaded to
a worker thread. Malformed content and content with no text both raise
:class:`TextExtractionError` — the ingestion pipeline turns that into a
``FAILED`` document with a recorded reason.
"""

from __future__ import annotations

import asyncio
from io import BytesIO

from app.domain.interfaces.extraction import TextExtractionError, TextExtractor
from app.platform.registry import register_provider

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


@register_provider(kind="extractor", name="text")
class PlainTextExtractor:
    """TXT/CSV: the content already is text (validated UTF-8 at upload)."""

    async def extract(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise TextExtractionError("File is not valid UTF-8 text") from exc


@register_provider(kind="extractor", name="pdf")
class PdfTextExtractor:
    """PDF via pypdf. Scanned/image-only PDFs yield no text and fail here;
    OCR support is a future provider behind the same interface."""

    async def extract(self, content: bytes) -> str:
        def _parse() -> str:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(content))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)

        try:
            text = await asyncio.to_thread(_parse)
        except TextExtractionError:
            raise
        except Exception as exc:
            raise TextExtractionError(f"Failed to parse PDF: {exc}") from exc
        if not text.strip():
            raise TextExtractionError(
                "PDF contains no extractable text (scanned documents need OCR)"
            )
        return text


@register_provider(kind="extractor", name="docx")
class DocxTextExtractor:
    """DOCX via python-docx: paragraphs plus table cell text."""

    async def extract(self, content: bytes) -> str:
        def _parse() -> str:
            import docx

            parsed = docx.Document(BytesIO(content))
            parts: list[str] = [
                paragraph.text for paragraph in parsed.paragraphs if paragraph.text.strip()
            ]
            for table in parsed.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append("\t".join(cells))
            return "\n".join(parts)

        try:
            text = await asyncio.to_thread(_parse)
        except Exception as exc:
            raise TextExtractionError(f"Failed to parse DOCX: {exc}") from exc
        if not text.strip():
            raise TextExtractionError("DOCX contains no extractable text")
        return text


@register_provider(kind="extractor", name="xlsx")
class XlsxTextExtractor:
    """XLSX via openpyxl: each row becomes a tab-joined line, sheets are titled."""

    async def extract(self, content: bytes) -> str:
        def _parse() -> str:
            from openpyxl import load_workbook

            workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
            try:
                parts: list[str] = []
                for sheet in workbook.worksheets:
                    rows: list[str] = []
                    for row in sheet.iter_rows(values_only=True):
                        values = ["" if cell is None else str(cell) for cell in row]
                        if any(value.strip() for value in values):
                            rows.append("\t".join(values))
                    if rows:
                        parts.append(f"# {sheet.title}\n" + "\n".join(rows))
                return "\n\n".join(parts)
            finally:
                workbook.close()

        try:
            text = await asyncio.to_thread(_parse)
        except Exception as exc:
            raise TextExtractionError(f"Failed to parse XLSX: {exc}") from exc
        if not text.strip():
            raise TextExtractionError("XLSX contains no extractable text")
        return text


_EXTRACTORS: dict[str, TextExtractor] = {
    "text/plain": PlainTextExtractor(),
    "text/csv": PlainTextExtractor(),
    "application/pdf": PdfTextExtractor(),
    _DOCX_MIME: DocxTextExtractor(),
    _XLSX_MIME: XlsxTextExtractor(),
}


def extractor_for(mime_type: str) -> TextExtractor:
    """Return the extractor for a canonical MIME type."""
    extractor = _EXTRACTORS.get(mime_type)
    if extractor is None:
        raise TextExtractionError(f"No text extractor for MIME type {mime_type!r}")
    return extractor
